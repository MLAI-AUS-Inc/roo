import io
import sys
from pathlib import Path
from types import SimpleNamespace

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from roo import linear_meeting_sources as sources
from roo import slack_client


def test_source_text_chunks_overlap_only_when_hard_splitting_long_paragraph():
    source = sources.ParsedSource(
        label="notes.pdf",
        text="A" * 12_000,
        kind="pdf",
    )

    chunks = sources.source_text_chunks(
        source,
        max_chars=5_000,
        hard_split_overlap_chars=300,
    )
    bodies = [chunk.split("\n", 1)[1] for chunk in chunks]

    assert [len(body) for body in bodies] == [5_000, 5_000, 2_600]
    assert bodies[0][-300:] == bodies[1][:300]
    assert bodies[1][-300:] == bodies[2][:300]


def test_get_thread_messages_preserves_file_metadata(monkeypatch):
    class FakeSlackClient:
        def conversations_replies(self, **kwargs):
            return {
                "ok": True,
                "messages": [
                    {
                        "user": "U1",
                        "text": "Attached the meeting notes.",
                        "ts": "1.1",
                        "files": [{"id": "F1", "name": "notes.pdf"}],
                    }
                ],
            }

    monkeypatch.setattr(slack_client, "get_slack_client", lambda: FakeSlackClient())

    messages = slack_client.get_thread_messages(channel="C1", thread_ts="1.1")

    assert messages[0]["files"] == [{"id": "F1", "name": "notes.pdf"}]


def test_get_thread_messages_follows_cursor_pagination(monkeypatch):
    requests = []

    class FakeSlackClient:
        def conversations_replies(self, **kwargs):
            requests.append(kwargs)
            if not kwargs.get("cursor"):
                return {
                    "ok": True,
                    "messages": [{"user": "U1", "text": "first", "ts": "1.1"}],
                    "response_metadata": {"next_cursor": "page-2"},
                }
            return {
                "ok": True,
                "messages": [{"user": "U2", "text": "latest", "ts": "1.2"}],
                "response_metadata": {"next_cursor": ""},
            }

    monkeypatch.setattr(slack_client, "get_slack_client", lambda: FakeSlackClient())

    messages = slack_client.get_thread_messages(channel="C1", thread_ts="1.1")

    assert [message["text"] for message in messages] == ["first", "latest"]
    assert requests == [
        {"channel": "C1", "ts": "1.1", "limit": 50},
        {"channel": "C1", "ts": "1.1", "limit": 50, "cursor": "page-2"},
    ]


def test_slack_file_info_and_download(monkeypatch):
    class FakeSlackClient:
        def files_info(self, file):
            assert file == "F1"
            return {
                "ok": True,
                "file": {
                    "id": "F1",
                    "name": "notes.txt",
                    "url_private_download": "https://files.slack.test/notes.txt",
                },
            }

    class FakeResponse:
        content = b"Sam will update the launch notes."

        def raise_for_status(self):
            return None

    request = {}

    def fake_get(url, headers, timeout):
        request.update({"url": url, "headers": headers, "timeout": timeout})
        return FakeResponse()

    monkeypatch.setattr(slack_client, "get_slack_client", lambda: FakeSlackClient())
    monkeypatch.setattr(slack_client, "get_settings", lambda: SimpleNamespace(SLACK_BOT_TOKEN="xoxb-test"))
    monkeypatch.setattr(slack_client.httpx, "get", fake_get)

    file = slack_client.get_file_info("F1")
    content = slack_client.download_file_bytes(file)

    assert file["name"] == "notes.txt"
    assert content == b"Sam will update the launch notes."
    assert request["headers"] == {"Authorization": "Bearer xoxb-test"}


def test_slack_file_download_reports_missing_files_read_scope(monkeypatch):
    class FakeSlackClient:
        def files_info(self, file):
            raise Exception("missing_scope needed=files:read")

    class FakeResponse:
        status_code = 302
        content = b""
        headers = {"location": "https://mlai-aus.slack.com/?redir=%2Ffiles-pri%2Ffile"}

        def raise_for_status(self):
            return None

    monkeypatch.setattr(slack_client, "get_slack_client", lambda: FakeSlackClient())
    monkeypatch.setattr(slack_client, "get_settings", lambda: SimpleNamespace(SLACK_BOT_TOKEN="xoxb-test"))
    monkeypatch.setattr(slack_client.httpx, "get", lambda *args, **kwargs: FakeResponse())

    with pytest.raises(RuntimeError, match="files:read"):
        slack_client.download_file_bytes(
            {
                "id": "F1",
                "name": "meeting.pdf",
                "url_private_download": "https://files.slack.test/meeting.pdf",
            }
        )


@pytest.mark.asyncio
async def test_parse_text_markdown_and_csv_sources(monkeypatch):
    downloads = {
        "notes.txt": b"Sam will update onboarding docs.",
        "notes.md": b"- Jane will send the launch summary.",
        "actions.csv": b"owner,task\nAlex,Review launch metrics",
    }

    monkeypatch.setattr(
        sources,
        "download_file_bytes",
        lambda file: downloads[file["name"]],
    )

    result = await sources.parse_linear_meeting_sources(
        text="send these files to Linear",
        params={},
        event_files=[
            {"id": "F1", "name": "notes.txt", "size": 10, "url_private_download": "https://x/1"},
            {"id": "F2", "name": "notes.md", "size": 10, "url_private_download": "https://x/2"},
            {"id": "F3", "name": "actions.csv", "size": 10, "url_private_download": "https://x/3"},
        ],
    )

    combined = result.combined_text()
    assert result.files_parsed == 3
    assert "Sam will update onboarding docs." in combined
    assert "Jane will send the launch summary." in combined
    assert "Review launch metrics" in combined


@pytest.mark.asyncio
async def test_parse_docx_source(monkeypatch):
    from docx import Document

    document = Document()
    document.add_paragraph("Taylor will draft the customer rollout checklist.")
    buffer = io.BytesIO()
    document.save(buffer)

    monkeypatch.setattr(sources, "download_file_bytes", lambda file: buffer.getvalue())

    result = await sources.parse_linear_meeting_sources(
        text="send this docx to Linear",
        params={},
        event_files=[
            {"id": "F1", "name": "notes.docx", "size": 10, "url_private_download": "https://x/docx"}
        ],
    )

    assert result.files_parsed == 1
    assert "Taylor will draft the customer rollout checklist." in result.combined_text()


@pytest.mark.asyncio
async def test_parse_text_pdf_source_with_mocked_pymupdf(monkeypatch):
    class FakePage:
        def get_text(self, mode):
            assert mode == "text"
            return "Sam will update the project plan with owners, milestones, risks, and launch dates."

    class FakeDocument:
        def __len__(self):
            return 1

        def load_page(self, index):
            assert index == 0
            return FakePage()

        def close(self):
            return None

    fake_fitz = SimpleNamespace(
        open=lambda stream, filetype: FakeDocument(),
        Matrix=lambda x, y: (x, y),
    )

    monkeypatch.setitem(sys.modules, "fitz", fake_fitz)
    monkeypatch.setattr(sources, "download_file_bytes", lambda file: b"%PDF-test")

    result = await sources.parse_linear_meeting_sources(
        text="send this pdf to Linear",
        params={},
        event_files=[
            {"id": "F1", "name": "meeting.pdf", "size": 10, "url_private_download": "https://x/pdf"}
        ],
    )

    assert result.files_parsed == 1
    assert "[meeting.pdf page 1]" in result.combined_text()
    assert "Sam will update the project plan" in result.combined_text()


@pytest.mark.asyncio
async def test_parse_image_source_uses_vision_parser(monkeypatch):
    calls = []

    async def image_parser(image_bytes, mime_type, label):
        calls.append((image_bytes, mime_type, label))
        return "Priya will follow up on the invoice list."

    monkeypatch.setattr(sources, "download_file_bytes", lambda file: b"image-bytes")

    result = await sources.parse_linear_meeting_sources(
        text="send this image to Linear",
        params={},
        event_files=[
            {
                "id": "F1",
                "name": "todos.png",
                "mimetype": "image/png",
                "size": 10,
                "url_private_download": "https://x/png",
            }
        ],
        image_parser=image_parser,
    )

    assert calls == [(b"image-bytes", "image/png", "todos.png")]
    assert "Priya will follow up on the invoice list." in result.combined_text()


@pytest.mark.asyncio
async def test_parse_sources_reports_unsupported_and_oversized_files(monkeypatch):
    downloaded = False

    def fake_download(file):
        nonlocal downloaded
        downloaded = True
        return b"unused"

    monkeypatch.setattr(sources, "download_file_bytes", fake_download)

    result = await sources.parse_linear_meeting_sources(
        text="send this file to Linear",
        params={},
        max_file_bytes=5,
        event_files=[
            {"id": "F1", "name": "legacy.doc", "size": 4, "url_private_download": "https://x/doc"},
            {"id": "F2", "name": "huge.pdf", "size": 6, "url_private_download": "https://x/pdf"},
        ],
    )

    assert result.files_skipped == 2
    assert not downloaded
    assert any("not supported" in warning for warning in result.warnings)
    assert any("larger than" in warning for warning in result.warnings)
