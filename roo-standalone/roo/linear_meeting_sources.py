from __future__ import annotations

import asyncio
import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Awaitable, Callable, Optional

from .slack_client import download_file_bytes, get_file_info


ImageParser = Callable[[bytes, str, str], Awaitable[str]]

TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCX_EXTENSIONS | PDF_EXTENSIONS | IMAGE_EXTENSIONS

DEFAULT_MAX_FILES = 5
DEFAULT_MAX_FILE_BYTES = 25 * 1024 * 1024
DEFAULT_MAX_PDF_PAGES = 75
DEFAULT_MAX_VISION_IMAGES = 10
PDF_TEXT_WORD_THRESHOLD = 8


@dataclass
class SlackSource:
    """Raw source discovered from Slack text or file metadata."""

    kind: str
    label: str
    text: str = ""
    file: Optional[dict[str, Any]] = None
    user: Optional[str] = None
    ts: Optional[str] = None


@dataclass
class ParsedSource:
    """Text extracted from one source for downstream action-item extraction."""

    label: str
    text: str
    kind: str
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class SourceParseResult:
    sources: list[ParsedSource] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    files_seen: int = 0
    files_parsed: int = 0
    files_skipped: int = 0

    def combined_text(self) -> str:
        parts = []
        for source in self.sources:
            text = source.text.strip()
            if text:
                parts.append(f"Source: {source.label}\n{text}")
        return "\n\n".join(parts).strip()


async def parse_linear_meeting_sources(
    *,
    text: str,
    params: dict[str, Any],
    thread_history: Optional[list[dict[str, Any]]] = None,
    event_files: Optional[list[dict[str, Any]]] = None,
    image_parser: Optional[ImageParser] = None,
    current_message_ts: Optional[str] = None,
    exclude_current_message: bool = False,
    max_files: int = DEFAULT_MAX_FILES,
    max_file_bytes: int = DEFAULT_MAX_FILE_BYTES,
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_vision_images: int = DEFAULT_MAX_VISION_IMAGES,
) -> SourceParseResult:
    """Parse Slack text and supported Slack files into text sources."""
    result = SourceParseResult()

    text_source = _build_thread_text_source(
        text,
        params,
        thread_history,
        current_message_ts=current_message_ts,
        exclude_current_message=exclude_current_message,
    )
    if text_source and text_source.text.strip():
        result.sources.append(
            ParsedSource(label=text_source.label, text=text_source.text, kind="slack_text")
        )

    files = _collect_slack_files(thread_history or [], event_files or [])
    result.files_seen = len(files)
    vision_budget = {"remaining": max_vision_images}

    for index, file in enumerate(files):
        if index >= max_files:
            result.files_skipped += 1
            result.warnings.append(
                f"Skipped `{_file_label(file)}` because only {max_files} files are processed per request."
            )
            continue

        try:
            full_file = await asyncio.to_thread(_resolve_file_metadata, file)
            label = _file_label(full_file)
            size = int(full_file.get("size") or 0)
            if size > max_file_bytes:
                result.files_skipped += 1
                result.warnings.append(
                    f"Skipped `{label}` because it is larger than {max_file_bytes // (1024 * 1024)} MB."
                )
                continue

            extension = _file_extension(full_file)
            if extension not in SUPPORTED_EXTENSIONS:
                result.files_skipped += 1
                result.warnings.append(
                    f"Skipped `{label}` because `{extension or 'unknown'}` files are not supported yet."
                )
                continue

            file_bytes = await asyncio.to_thread(download_file_bytes, full_file)
            parsed = await _parse_file_bytes(
                file=full_file,
                file_bytes=file_bytes,
                image_parser=image_parser,
                vision_budget=vision_budget,
                max_pdf_pages=max_pdf_pages,
            )
            if parsed:
                result.sources.extend(parsed)
                result.files_parsed += 1
            else:
                result.files_skipped += 1
                result.warnings.append(f"No readable text was found in `{label}`.")
        except Exception as exc:
            result.files_skipped += 1
            result.warnings.append(
                f"Could not parse `{_file_label(file)}`: {exc.__class__.__name__}: {exc}"
            )

    return result


def _build_thread_text_source(
    text: str,
    params: dict[str, Any],
    thread_history: Optional[list[dict[str, Any]]],
    *,
    current_message_ts: Optional[str] = None,
    exclude_current_message: bool = False,
) -> Optional[SlackSource]:
    explicit = str(params.get("transcript") or "").strip()
    parts: list[str] = []
    if explicit:
        parts.append(explicit)

    current_ts = str(current_message_ts or "").strip()
    for message in thread_history or []:
        if message.get("is_bot") or message.get("bot_id"):
            continue
        if exclude_current_message and current_ts and str(message.get("ts") or "").strip() == current_ts:
            continue
        message_text = str(message.get("text") or "").strip()
        if not message_text:
            continue
        speaker = _slack_speaker_label(message)
        source_prefix = _slack_message_source_prefix(message)
        parts.append(f"{source_prefix}{speaker}: {message_text}")

    clean_text = str(text or "").strip()
    if not exclude_current_message and clean_text and all(clean_text not in part for part in parts):
        parts.append(clean_text)

    if not parts:
        return None
    return SlackSource(kind="text", label="Slack thread", text="\n".join(dict.fromkeys(parts)))


def _slack_speaker_label(message: dict[str, Any]) -> str:
    user_id = str(message.get("user") or "").strip()
    display_name = str(message.get("display_name") or "").strip()
    email = str(message.get("email") or "").strip()
    identity_parts: list[str] = []
    if user_id:
        identity_parts.append(f"<@{user_id}>")
    if email:
        identity_parts.append(email)
    if display_name and identity_parts:
        return f"{display_name} ({', '.join(identity_parts)})"
    return display_name or (f"<@{user_id}>" if user_id else "user")


def _slack_message_source_prefix(message: dict[str, Any]) -> str:
    metadata: list[str] = []
    ts = str(message.get("ts") or "").strip()
    local_datetime = str(message.get("local_datetime") or "").strip()
    if ts:
        metadata.append(f"ts={ts}")
    if local_datetime:
        metadata.append(f"local={local_datetime}")
    return f"[Slack message {' '.join(metadata)}] " if metadata else ""


def _collect_slack_files(
    thread_history: list[dict[str, Any]],
    event_files: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    files: list[dict[str, Any]] = []
    seen: set[str] = set()

    def add_file(file: dict[str, Any]) -> None:
        if not isinstance(file, dict):
            return
        key = (
            str(file.get("id") or "").strip()
            or str(file.get("url_private_download") or file.get("url_private") or "").strip()
            or str(file.get("name") or file.get("title") or "").strip()
        )
        if not key or key in seen:
            return
        seen.add(key)
        files.append(file)

    for file in event_files:
        add_file(file)
    for message in thread_history:
        for file in message.get("files") or []:
            add_file(file)

    return files


def _resolve_file_metadata(file: dict[str, Any]) -> dict[str, Any]:
    file_id = str(file.get("id") or "").strip()
    needs_info = (
        file.get("file_access") == "check_file_info"
        or not (file.get("url_private_download") or file.get("url_private"))
    )
    if needs_info and file_id:
        return get_file_info(file_id)
    return dict(file)


async def _parse_file_bytes(
    *,
    file: dict[str, Any],
    file_bytes: bytes,
    image_parser: Optional[ImageParser],
    vision_budget: dict[str, int],
    max_pdf_pages: int,
) -> list[ParsedSource]:
    extension = _file_extension(file)
    label = _file_label(file)

    if extension in TEXT_EXTENSIONS:
        return [ParsedSource(label=label, text=_decode_text(file_bytes), kind=extension.lstrip("."))]

    if extension in DOCX_EXTENSIONS:
        return [ParsedSource(label=label, text=_parse_docx(file_bytes), kind="docx")]

    if extension in PDF_EXTENSIONS:
        return await _parse_pdf(
            label=label,
            file_bytes=file_bytes,
            image_parser=image_parser,
            vision_budget=vision_budget,
            max_pdf_pages=max_pdf_pages,
        )

    if extension in IMAGE_EXTENSIONS:
        if extension == ".gif" and _is_animated_gif(file_bytes):
            raise RuntimeError("Animated GIF files are not supported yet")
        text = await _parse_image(
            label=label,
            file_bytes=file_bytes,
            mime_type=_file_mimetype(file),
            image_parser=image_parser,
            vision_budget=vision_budget,
        )
        return [ParsedSource(label=label, text=text, kind="image")] if text.strip() else []

    return []


def _decode_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace").strip()


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to parse DOCX files") from exc

    document = Document(io.BytesIO(file_bytes))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts).strip()


async def _parse_pdf(
    *,
    label: str,
    file_bytes: bytes,
    image_parser: Optional[ImageParser],
    vision_budget: dict[str, int],
    max_pdf_pages: int,
) -> list[ParsedSource]:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PyMuPDF is required to parse PDF files") from exc

    document = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        page_count = min(len(document), max_pdf_pages)
        parts: list[str] = []

        for page_index in range(page_count):
            page = document.load_page(page_index)
            page_label = f"{label} page {page_index + 1}"
            page_text = (page.get_text("text") or "").strip()
            if len(page_text.split()) >= PDF_TEXT_WORD_THRESHOLD:
                parts.append(f"[{page_label}]\n{page_text}")
                continue

            if image_parser and vision_budget["remaining"] > 0:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image_bytes = pixmap.tobytes("png")
                image_text = await _parse_image(
                    label=page_label,
                    file_bytes=image_bytes,
                    mime_type="image/png",
                    image_parser=image_parser,
                    vision_budget=vision_budget,
                )
                if image_text.strip():
                    parts.append(f"[{page_label}]\n{image_text.strip()}")

        if len(document) > max_pdf_pages:
            parts.append(f"[Parser note] Remaining pages after page {max_pdf_pages} were not processed.")
    finally:
        close = getattr(document, "close", None)
        if callable(close):
            close()

    text = "\n\n".join(parts).strip()
    return [ParsedSource(label=label, text=text, kind="pdf")] if text else []


async def _parse_image(
    *,
    label: str,
    file_bytes: bytes,
    mime_type: str,
    image_parser: Optional[ImageParser],
    vision_budget: dict[str, int],
) -> str:
    if not image_parser:
        raise RuntimeError("Image parsing is unavailable because OPENAI_API_KEY is not configured")
    if vision_budget["remaining"] <= 0:
        raise RuntimeError("Image parsing limit reached for this request")
    vision_budget["remaining"] -= 1
    return (await image_parser(file_bytes, mime_type, label)).strip()


def _file_label(file: dict[str, Any]) -> str:
    return str(
        file.get("name")
        or file.get("title")
        or file.get("id")
        or "Slack file"
    ).strip()


def _file_extension(file: dict[str, Any]) -> str:
    name = _file_label(file)
    extension = Path(name).suffix.lower()
    if extension:
        return extension

    filetype = str(file.get("filetype") or "").lower().strip(".")
    if filetype:
        return f".{filetype}"

    mimetype = _file_mimetype(file)
    if mimetype == "application/pdf":
        return ".pdf"
    if mimetype in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        return ".docx" if "openxmlformats" in mimetype else ".doc"
    if mimetype.startswith("image/"):
        return f".{mimetype.split('/', 1)[1].replace('jpeg', 'jpg')}"
    if mimetype.startswith("text/"):
        return ".txt"
    return ""


def _file_mimetype(file: dict[str, Any]) -> str:
    mimetype = str(file.get("mimetype") or "").strip().lower()
    if mimetype:
        return mimetype

    extension = Path(_file_label(file)).suffix.lower()
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".pdf": "application/pdf",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".csv": "text/csv",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }.get(extension, "application/octet-stream")


def _is_animated_gif(file_bytes: bytes) -> bool:
    if not file_bytes.startswith((b"GIF87a", b"GIF89a")):
        return False
    return file_bytes.count(b"\x21\xf9\x04") > 1


def source_text_chunks(
    source: ParsedSource,
    max_chars: int = 10000,
    *,
    hard_split_overlap_chars: int = 0,
) -> list[str]:
    """Split source text into chunks without dropping source attribution."""
    if max_chars <= 0:
        raise ValueError("max_chars must be positive")
    if hard_split_overlap_chars < 0 or hard_split_overlap_chars >= max_chars:
        raise ValueError("hard_split_overlap_chars must be between 0 and max_chars")

    text = source.text.strip()
    if len(text) <= max_chars:
        return [f"Source: {source.label}\n{text}"] if text else []

    chunks: list[str] = []
    paragraphs = re.split(r"\n{2,}", text)
    current = ""
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if len(paragraph) > max_chars:
            if current:
                chunks.append(f"Source: {source.label}\n{current}")
                current = ""
            step = max_chars - hard_split_overlap_chars
            for start in range(0, len(paragraph), step):
                chunk_text = paragraph[start:start + max_chars]
                chunks.append(f"Source: {source.label}\n{chunk_text}")
                if start + max_chars >= len(paragraph):
                    break
            continue
        next_value = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(next_value) > max_chars and current:
            chunks.append(f"Source: {source.label}\n{current}")
            current = paragraph
        else:
            current = next_value
    if current:
        chunks.append(f"Source: {source.label}\n{current}")
    return chunks
